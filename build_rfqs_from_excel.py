from __future__ import annotations
import os, json, pathlib, textwrap, re, shutil
from typing import List, Dict, Tuple
import tempfile, time
import pandas as pd
from tqdm import tqdm
from dotenv import load_dotenv
import openai
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import qn, nsdecls

# ── ENV / KEY ───────────────────────────────────────────────
load_dotenv()
openai.api_key = os.getenv("OPENAI_API_KEY") or ""
if not openai.api_key:
    raise RuntimeError("OPENAI_API_KEY not set")

# ── PATHS & CONSTANTS ───────────────────────────────────────
EXCEL_PATH   = pathlib.Path("C:/Users/PRANAY-RES/OneDrive - Renewable Energy Systems Limited/RES/Capital Equipment 2025/Capital_Equipment.xlsx")
OLD_PATH   = pathlib.Path("data/old/data.xlsx")
TEMPLATE   = pathlib.Path("templates/u1.docx")
OUT_ROOT     = pathlib.Path("C:/Users/PRANAY-RES/OneDrive - Renewable Energy Systems Limited/RES/Capital Equipment 2025")
MODEL      = "gpt-4o-mini"
TEMPERATURE = 0.2
OUT_ROOT.mkdir(exist_ok=True, parents=True)

# ── 1. Load current & snapshot Excel data ───────────────────
def load_sheet(path: pathlib.Path) -> pd.DataFrame:
    """
    Safely read sheet 'cp_list' from *path*.

    • Copies the file to a temp name first (avoids OneDrive/Excel locks).
    • If the sheet is missing, returns an empty DataFrame.
    • Ensures 'cid' is int, NaNs in Item/Description become empty strings.
    """
    cols = ["cid", "Item", "Priority", "Assignee", "Description"]

    if not path.exists():
        return pd.DataFrame(columns=cols)

    # ---------- copy to temp to bypass locks ----------
    with tempfile.NamedTemporaryFile(suffix=".xlsx", delete=False) as tmp:
        tmp_path = pathlib.Path(tmp.name)
    shutil.copy2(path, tmp_path)

    try:
        try:
            df = pd.read_excel(tmp_path, sheet_name="cp_list", header=2).iloc[1:]
        except ValueError:  # sheet missing
            return pd.DataFrame(columns=cols)
    finally:
        tmp_path.unlink(missing_ok=True)

    df = df.rename(
        columns={
            "Unnamed: 1": "cid",
            "Unnamed: 2": "Item",
            "Unnamed: 3": "Priority",
            "Unnamed: 4": "Assignee",
            "Unnamed: 5": "Description",
        }
    )[cols]

    df["cid"] = pd.to_numeric(df["cid"], errors="coerce").astype("Int64")
    df = df.dropna(subset=["cid", "Item"]).fillna("")
    df["cid"] = df["cid"].astype(int)
    return df.reset_index(drop=True)


current_df = load_sheet(EXCEL_PATH)
old_df     = load_sheet(OLD_PATH)

# build lookup dicts: cid -> (item, desc)
def make_lookup(df: pd.DataFrame) -> Dict[int, Tuple[str,str]]:
    lut={}
    for _,row in df.iterrows():
        lut[int(row["cid"])] = (str(row["Item"]).strip(), str(row.get("Description","")).strip())
    return lut

cur_lut = make_lookup(current_df)
old_lut = make_lookup(old_df)

# determine which cids need (re)generation
to_process=[]
for cid, (item, desc) in cur_lut.items():
    if cid not in old_lut:
        to_process.append(cid)                 # new equipment
    else:
        old_item, old_desc = old_lut[cid]
        if item!=old_item or desc!=old_desc:   # modified
            to_process.append(cid)

if not to_process:
    print("✓ No new or modified equipment rows – nothing to generate.")
    exit()

# ── 2. GPT prompt templates (intro, scope, tech, docs) ─────
SYSTEM_PROMPT = textwrap.dedent("""\
You are a senior procurement engineer at Renewable Energy Systems Limited
(AS9100D, lithium‑thermal battery manufacturer). Draft concise RFQ sections
aligned with MIL‑STD‑810H, MIL‑STD‑1580, ASTM, IEC, etc.

Return **valid JSON only** with keys:
  introduction, scope, tech_table, docs_required

• tech_table must contain **at least 8** relevant {parameter, requirement} pairs.
""")

USER_TMPL = (
    "Draft the RFQ sections for equipment '{item}'. "
    "{desc_clause}"
    "Return JSON only."
)

def gpt_sections(item:str, desc:str|None)->dict:
    clause = f"Include these user‑specified criteria: {desc}. " if desc else ""
    prompt = USER_TMPL.format(item=item, desc_clause=clause)
    resp = openai.chat.completions.create(
        model=MODEL, temperature=TEMPERATURE,
        response_format={"type":"json_object"},
        messages=[{"role":"system","content":SYSTEM_PROMPT},
                  {"role":"user","content":prompt}]
    )
    return json.loads(resp.choices[0].message.content)

# ── 3. DOCX helpers (style, borders, bullets) ───────────────
def style_body(par):
    for r in par.runs:
        r.font.bold=False; r.font.underline=False; r.font.size=Pt(11)

def bullet(doc, entry, bullet_ok):
    if isinstance(entry,(list,tuple)):
        for e in entry: bullet(doc,e,bullet_ok); return
    p = doc.add_paragraph(entry,style="List Bullet") if bullet_ok else doc.add_paragraph(f"• {entry}")
    style_body(p)

def box(cell):
    tcPr=cell._tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:tcBorders")): tcPr.remove(old)
    tcPr.append(parse_xml(
        r'<w:tcBorders %s>'
        r'<w:top w:val="single" w:sz="4" w:color="000000"/>'
        r'<w:left w:val="single" w:sz="4" w:color="000000"/>'
        r'<w:bottom w:val="single" w:sz="4" w:color="000000"/>'
        r'<w:right w:val="single" w:sz="4" w:color="000000"/>'
        r'</w:tcBorders>' % nsdecls('w')))

COMM_ROWS=[
    ("Quotation Validity","Minimum 90 days"),
    ("Delivery Time","Specify lead time"),
    ("Pricing Terms","Provide EXW, FOB, and CIF pricing (as applicable)"),
    ("Payment Terms","To be negotiated"),
    ("Installation and Training","Specify installation and operator training charges if applicable"),
    ("After‑Sales Support","Provide details of service support, spares availability, and annual maintenance contracts"),
]

def build_doc(sec:dict,item:str,cid:int,bullet_ok:bool)->pathlib.Path:
    doc=Document(TEMPLATE)
    if doc.paragraphs and not doc.paragraphs[0].text.strip():
        p=doc.paragraphs[0]._element; p.getparent().remove(p)

    title=doc.add_paragraph(f"RFQ for {item}")
    title.alignment=WD_ALIGN_PARAGRAPH.CENTER
    title.style.font.size=Pt(16); title.style.font.bold=True
    title.paragraph_format.space_before=Pt(0); title.paragraph_format.space_after=Pt(0)

    def H(n,t):
        h=doc.add_heading(f"{n}. {t}",level=2)
        for r in h.runs: r.font.size=Pt(14); r.font.bold=True

    H(1,"Introduction"); style_body(doc.add_paragraph(sec["introduction"]))
    H(2,"Scope of Supply"); style_body(doc.add_paragraph(sec["scope"]))

    H(3,"Technical Requirements")
    tbl=doc.add_table(rows=1,cols=2); tbl.alignment=WD_TABLE_ALIGNMENT.CENTER
    hdr=tbl.rows[0].cells; hdr[0].text="Parameter"; hdr[1].text="Requirement"
    for c in hdr: r=c.paragraphs[0].runs[0]; r.font.bold=True; r.font.size=Pt(12); box(c)
    for row in sec["tech_table"]:
        cells=tbl.add_row().cells
        cells[0].text,row_param = row["parameter"], row["requirement"]
        cells[1].text = row_param
        for c in cells: style_body(c.paragraphs[0]); box(c)

    H(4,"Commercial Requirements")
    ct=doc.add_table(rows=1,cols=2); ct.alignment=WD_TABLE_ALIGNMENT.CENTER
    ch=ct.rows[0].cells; ch[0].text="Parameter"; ch[1].text="Requirement"
    for c in ch: r=c.paragraphs[0].runs[0]; r.font.bold=True; r.font.size=Pt(12); box(c)
    for p,rq in COMM_ROWS:
        cells=ct.add_row().cells; cells[0].text, cells[1].text = p, rq
        for c in cells: style_body(c.paragraphs[0]); box(c)

    H(5,"Documentation Requirements")
    bullet(doc, sec["docs_required"], bullet_ok)

    H(6,"Submission Guidelines")
    style_body(doc.add_paragraph(
        f"Please submit your quotations via email with the subject line:\n"
        f"“Quotation for {item} - RESL”."))
    bullet(doc,"Contact Person: P. Pranay Kiran, A. Sai Nithin",bullet_ok)
    bullet(doc,"Email: Designengineer.pranay@resindia.co.in, engineer.resl1@resindia.co.in",bullet_ok)

    H(7,"Confidentiality Clause")
    style_body(doc.add_paragraph(
        "All quotations and related documents submitted in response to this RFQ "
        "will be treated as confidential and used solely for evaluation purposes."))

    safe=re.sub(r"[^A-Za-z0-9_]+","_",item.strip())
    rfq_dir=(OUT_ROOT/safe/"RFQ"); rfq_dir.mkdir(parents=True, exist_ok=True)
    out=rfq_dir/f"RFQ_{cid:02d}_{safe}.docx"; doc.save(out); return out

# bullet style availability
try: Document(TEMPLATE).styles["List Bullet"]; BULLET_OK=True
except KeyError: BULLET_OK=False

# ── 5. Generate only new/changed rows ──────────────────────
created=[]
rows_to_process=current_df[current_df["cid"].isin(to_process)]

for _,row in tqdm(rows_to_process.iterrows(), total=len(rows_to_process),
                  desc="Generating RFQs"):
    cid=int(row["cid"])
    item=str(row["Item"]).strip()
    desc=row.get("Description","")
    try:
        sections=gpt_sections(item, desc)
    except Exception as e:
        print(f"\n⚠️ GPT error on cid {cid} '{item}': {e}"); continue
    try:
        created.append(build_doc(sections,item,cid,BULLET_OK))
    except Exception as e:
        print(f"\n⚠️ DOCX error on cid {cid} '{item}': {e}")

print(f"\n✓ Generated/updated {len(created)} RFQ file(s).")

# ── 6. Update snapshot ─────────────────────────────────────
OLD_PATH.parent.mkdir(parents=True, exist_ok=True)
shutil.copy2(EXCEL_PATH, OLD_PATH)
print(f"✓ Snapshot updated → {OLD_PATH}")
