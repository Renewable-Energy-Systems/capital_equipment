"""
build_proposals_ai.py – v2.0
============================
AI-drafted Capital-Equipment Proposals

• Reads sheet 'cp_list' in Capital_Equipment.xlsx
• Feeds *all* useful columns (Item, Description, Impact Analysis,
  Ref./NC no., Proposed On/By, Org Details) to GPT-4o.
• GPT returns JSON for every narrative section in the template.
• Fills CP Proposal Document.docx placeholders and saves to
    <OUT_ROOT>/<Item>/Proposal/Proposal_<cid>_<Item>.docx
"""

from __future__ import annotations
import os, json, pathlib, re, tempfile, shutil, sys
from typing import Dict

import pandas as pd
from tqdm import tqdm
from dotenv import load_dotenv
import openai
from docx import Document
from docx.shared import Pt

# ── ENV / API KEY ────────────────────────────────────────────
load_dotenv()
openai.api_key = os.getenv("OPENAI_API_KEY", "")
if not openai.api_key:
    sys.exit("❌  OPENAI_API_KEY not set")

MODEL       = "gpt-4o-mini"   # or "gpt-4o"
TEMPERATURE = 0.25

# ── FILE LOCATIONS ──────────────────────────────────────────
EXCEL_PATH = pathlib.Path(
    r"C:/Users/PRANAY-RES/OneDrive - Renewable Energy Systems Limited/RES/Capital Equipment 2025/Capital_Equipment.xlsx"
)
TEMPLATE   = pathlib.Path("templates/CP Proposal Document.docx")
OUT_ROOT   = pathlib.Path(
    r"C:/Users/PRANAY-RES/OneDrive - Renewable Energy Systems Limited/RES/Capital Equipment 2025"
)
SHEET_NAME = "cp_list"
OUT_ROOT.mkdir(parents=True, exist_ok=True)

# ── Template placeholders (exact strings) ───────────────────
PH = {
    "intro":    "Provide a concise overview of the capital equipment requirement, its strategic alignment with organizational goals, and the anticipated benefits of acquiring the equipment.",
    "reason":   "Explain the operational, quality, compliance, or capacity drivers that necessitate the acquisition of this equipment.",
    "benefits": "Summarize the impact on safety, quality, productivity, cost, compliance, and customer satisfaction.",
    "oper":     "Outline the estimated running costs, service intervals, and required consumables.",
    "roi":      "Present a quantitative ROI analysis, including payback period, net present value (NPV), or internal rate of return (IRR) as applicable.",
    "timeline": "Provide key milestones from the purchase order to commissioning and operator training.",
    "risks":    "Identify potential risks and proposed mitigation actions.",
    "concl":    "Reiterate the justification and formally request approval for the capital expenditure.",
}

# ── Helper: sanitise file/dir names ─────────────────────────
def safe(txt: str) -> str:
    return re.sub(r"[^A-Za-z0-9_]+", "_", txt.strip())

# ── Read sheet robustly (detects header row) ───────────────
def read_cp_list(xlsx: pathlib.Path) -> pd.DataFrame:
    if not xlsx.exists():
        sys.exit(f"❌ Workbook not found: {xlsx}")
    # copy to temp (avoids OneDrive lock)
    with tempfile.NamedTemporaryFile(suffix=".xlsx", delete=False) as tmp:
        tmp_path = pathlib.Path(tmp.name)
    shutil.copy2(xlsx, tmp_path)

    try:
        raw = pd.read_excel(tmp_path, sheet_name=SHEET_NAME, header=None)
    finally:
        tmp_path.unlink(missing_ok=True)

    # detect header row
    hdr = next(
        i
        for i, row in raw.iterrows()
        if {"cid", "item"} & {str(c).strip().lower() for c in row}
    )
    raw.columns = raw.iloc[hdr]
    df = raw.iloc[hdr + 1 :].fillna("").reset_index(drop=True)

    canon = {}
    for col in df.columns:
        c = str(col).strip().lower()
        if c in ("cid", "id"):
            canon[col] = "cid"
        elif c in ("item", "equipment", "equipment name"):
            canon[col] = "item"
        elif "description" in c:
            canon[col] = "description"
        elif "impact" in c:
            canon[col] = "impact"
        elif "ref" in c or "nc" in c:
            canon[col] = "ref_no"
        elif "proposed on" in c or "date" in c:
            canon[col] = "proposed_on"
        elif "proposed by" in c:
            canon[col] = "proposed_by"
        elif "organization" in c:
            canon[col] = "org_details"

    df = df.rename(columns=canon)
    if {"cid", "item"} - set(df.columns):
        sys.exit("❌ Columns 'cid' and 'item' are required but missing.")
    df["cid"] = pd.to_numeric(df["cid"], errors="coerce").astype("Int64")
    df = df.dropna(subset=["cid", "item"]).reset_index(drop=True)
    df["cid"] = df["cid"].astype(int)
    return df

# ── GPT generation ─────────────────────────────────────────
SYSTEM_PROMPT = (
    "You are the Capital-Expenditure engineering assistant for Renewable Energy "
    "Systems Limited (AS9100D, lithium-thermal battery maker). "
    "Write succinct, professional proposal sections. "
    "You may invent reasonable industry-average figures for costs, ROI, "
    "maintenance, etc. when not provided.\n"
    "Return **JSON only** with keys:\n"
    " introduction, reason, benefits, operating, roi, timeline, risks, conclusion."
)

def gpt_sections(ctx: Dict[str, str]) -> Dict[str, str]:
    # Build user message summarising available columns
    user = (
        f"Equipment name: {ctx['item']}\n"
        f"Reference / NC no.: {ctx.get('ref_no','N/A')}\n"
        f"Description (user notes/specs): {ctx.get('description','N/A')}\n"
        f"Impact analysis (user notes): {ctx.get('impact','N/A')}\n"
        f"Proposed On: {ctx.get('proposed_on','N/A')}\n"
        f"Proposed By: {ctx.get('proposed_by','N/A')}\n"
        f"Organisation details: {ctx.get('org_details','N/A')}\n\n"
        "Using this context draft the proposal sections."
    )
    resp = openai.chat.completions.create(
        model=MODEL,
        temperature=TEMPERATURE,
        response_format={"type": "json_object"},
        messages=[{"role":"system","content":SYSTEM_PROMPT},
                  {"role":"user","content":user}],
    )
    return json.loads(resp.choices[0].message.content)

# ── Fill Word template placeholders ────────────────────────
def fill_template(row: pd.Series, ai: Dict[str, str]) -> Document:
    doc = Document(TEMPLATE)
    
    def _to_text(v) -> str:          # ← NEW
        if isinstance(v, list):
            return "\n".join(map(str, v))
        return str(v)

    replace_map = {
        PH["intro"]:  _to_text(ai["introduction"]),
        PH["reason"]: _to_text(ai["reason"]),
        PH["benefits"]: _to_text(ai["benefits"]),
        PH["oper"]:  _to_text(ai["operating"]),
        PH["roi"]:   _to_text(ai["roi"]),
        PH["timeline"]: _to_text(ai["timeline"]),
        PH["risks"]: _to_text(ai["risks"]),
        PH["concl"]: _to_text(ai["conclusion"]),
    }
    for para in doc.paragraphs:
        txt = para.text
        for old, new in replace_map.items():
            if old in txt:
                para.text = txt.replace(old, new)
                txt = para.text
        # inline form fields
        if "[Enter applicable NC or reference number]" in txt and row.get("ref_no",""):
            para.text = txt.replace(
                "[Enter applicable NC or reference number]", str(row["ref_no"])
            )
        if "Proposed On:" in txt and row.get("proposed_on",""):
            para.text = txt.replace("____________________", str(row["proposed_on"]))
        if "Proposed By:" in txt and row.get("proposed_by",""):
            para.text = txt.replace("____________________", str(row["proposed_by"]))
        if "Organization Details:" in txt and row.get("org_details",""):
            para.text = txt.replace(
                "_______________________________________________", str(row["org_details"])
            )

        # body font
        for r in para.runs:
            r.font.size = Pt(11)
            r.font.bold = False

    # style title
    for p in doc.paragraphs:
        if p.text.lower().startswith("proposal for"):
            p.alignment = 1
            for r in p.runs:
                r.font.size = Pt(16); r.font.bold = True
            break
    return doc

# ── MAIN ────────────────────────────────────────────────────
df = read_cp_list(EXCEL_PATH)

created = 0
for _, row in tqdm(df.iterrows(), total=len(df), desc="AI Proposals"):
    ctx = {k: str(row.get(k,"")).strip() for k in row.index}
    try:
        ai = gpt_sections(ctx)
    except Exception as e:
        print(f"\n⚠️ GPT error cid {row['cid']} '{row['item']}': {e}")
        continue

    doc = fill_template(row, ai)

    out_dir = OUT_ROOT / safe(row["item"]) / "Proposal"
    out_dir.mkdir(parents=True, exist_ok=True)
    out_file = out_dir / f"Proposal_{row['cid']:02d}_{safe(row['item'])}.docx"
    doc.save(out_file)
    created += 1

print(f"\n✓ {created} proposal document(s) written under '{OUT_ROOT}'.")
